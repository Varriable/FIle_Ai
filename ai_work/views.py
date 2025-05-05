from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import UploadedFile
from .forms import UploadFileForm, AskForm
import os
import requests
from PyPDF2 import PdfReader
from docx import Document


def register(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('file_list')
    else:
        form = UserCreationForm()
    return render(request, 'registration/register.html', {'form': form})


    # def extract_text(file_path):
    # ext = os.path.splitext(file_path)[1].lower()
    # if ext == '.txt':
    #     with open(file_path, 'r', encoding='utf-8') as f:
    #         return f.read()
    # elif ext == '.pdf':
    #     import PyPDF2
    #     with open(file_path, 'rb') as f:
    #         reader = PyPDF2.PdfReader(f)
    #         return "\n".join([page.extract_text() for page in reader.pages])
    # elif ext == '.docx':
    #     import docx
    #     doc = docx.Document(file_path)
    #     return "\n".join([p.text for p in doc.paragraphs])
    # return "Unsupported file format."
def extract_file_text(file_path):
    """Extract text from TXT, PDF, or DOCX files."""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif ext == ".pdf":
            reader = PdfReader(file_path)
            return " ".join([page.extract_text() for page in reader.pages if page.extract_text()])

        elif ext == ".docx":
            doc = Document(file_path)
            return " ".join([para.text for para in doc.paragraphs])

        else:
            return "Unsupported file type."

    except Exception as e:
        return f"Error reading file: {e}"

def ask_ai(prompt, context):
    import requests
    response = requests.post(
        "http://localhost:1234/v1/completions",
        json={
            "prompt": f"Context:\n{context}\n\nQuestion: {prompt}\nAnswer:",
            "max_tokens": 300,
            "temperature": 0.7
        }
    )
    return response.json()['choices'][0]['text'].strip()

@login_required
def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = form.save(commit=False)
            uploaded_file.user = request.user
            uploaded_file.save()
            return redirect('file_list')
    else:
        form = UploadFileForm()
    return render(request, 'ai_work/upload.html', {'form': form})

@login_required
def file_list(request):
    files = UploadedFile.objects.filter(user=request.user)
    return render(request, 'ai_work/list.html', {'files': files})

@login_required
# def file_detail(request, pk):
#     file = get_object_or_404(UploadedFile, pk=pk, user=request.user)
#     form = AskForm()
#     result = None
#     context_text = extract_text(file.file.path)

#     if request.method == 'POST':
#         form = AskForm(request.POST)
#         if form.is_valid():
#             prompt = form.cleaned_data.get('prompt')
#             if 'summarise' in request.POST:
#                 result = ask_ai("Summarise this document", context_text)
#             elif prompt:
#                 result = ask_ai(prompt, context_text)

#     return render(request, 'ai_work/detail.html', {
#         'file': file,
#         'form': form,
#         'result': result,
#     })
@login_required
def file_detail(request, pk):
    file = get_object_or_404(UploadedFile, pk=pk, user=request.user)
    ai_response = None

    if request.method == "POST":
        action = request.POST.get("action")
        question = request.POST.get("question", "").strip()

        file_path = file.file.path
        file_text = extract_file_text(file_path)

        if not file_text:
            ai_response = "Could not extract text from the file."
        else:
            if action == "ask" and question:
                prompt = f"Document:\n{file_text}\n\nQuestion: {question}\nAnswer:"
                ai_response = query_lm_studio(prompt)

            elif action == "summarize":
                prompt = f"Please summarize the following document:\n\n{file_text}\n\nSummary:"
                ai_response = query_lm_studio(prompt)

    return render(request, "ai_work/detail.html", {
        "file": file,
        "ai_response": ai_response
    })

def query_lm_studio(prompt):
    """Send a prompt to the locally hosted LM Studio model."""
    try:
        response = requests.post(
            "http://localhost:1234/v1/completions",
            headers={"Content-Type": "application/json"},
            json={
                "prompt": prompt,
                "max_tokens": 1024,
                "temperature": 0.7,
                "stop": None,
            }
        )
        data = response.json()
        return data.get("choices", [{}])[0].get("text", "").strip()
    except Exception as e:
        return f"Error connecting to AI: {e}"





    




